import os
from docx import Document
from PIL import Image

def clean_cell_text(text):
    """Cleans OCR text in a table cell to avoid spurious short fragments."""
    if not text:
        return ""
    stripped = text.strip()
    # Ignore very short text fragments (e.g. 'PO') which may be noise
    if len(stripped) <= 2:
        return ""
    return stripped

def write_ocr_output_to_docx(ocr_data, output_path):
    """
    Writes OCR extracted data to a DOCX file.

    Parameters:
        ocr_data (dict): Contains 'top_text' (str) and 'table' (list of lists or None).
        output_path (str): Path to save the DOCX file.
    """
    doc = Document()

    # Write top text paragraph(s)
    if ocr_data.get("top_text"):
        doc.add_paragraph(ocr_data["top_text"])

    # Write table if exists
    if ocr_data.get("table"):
        rows = ocr_data["table"]
        max_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cleaned_text = clean_cell_text(cell_text)
                table.cell(i, j).text = cleaned_text

    doc.save(output_path)
    print(f"Saved: {output_path}")

def batch_ocr_to_docx(input_folder, output_folder, ocr_function):
    """
    Process all PNG images in input_folder with ocr_function and write DOCX outputs.

    Parameters:
        input_folder (str): Folder containing PNG images.
        output_folder (str): Folder to save DOCX files.
        ocr_function (callable): Function that takes a PIL.Image and returns OCR data dict.
    """
    os.makedirs(output_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        if filename.lower().endswith(".png"):
            input_path = os.path.join(input_folder, filename)
            image = Image.open(input_path).convert("RGB")

            ocr_data = ocr_function(image)

            output_filename = filename.rsplit(".", 1)[0] + ".docx"
            output_path = os.path.join(output_folder, output_filename)

            write_ocr_output_to_docx(ocr_data, output_path)
