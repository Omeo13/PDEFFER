import os
import re
from docx import Document
from PIL import Image

def clean_cell_text(text):
    """Cleans OCR text in a table cell, aggressively filtering out likely hallucinations."""
    if not text:
        return ""

    stripped = text.strip()
    cleaned = re.sub(r"[^\w\s]", "", stripped)

    if len(cleaned) <= 2:
        return ""

    blacklist = {"po", "p0", "l0", "lo", "oo", "xx", "na", "n/a"}
    if cleaned.lower() in blacklist:
        return ""

    if not re.search(r"[a-zA-Z]", cleaned):
        return ""

    return cleaned

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import os

def write_ocr_output_to_docx(ocr_data, output_path):
    """
    Writes OCR extracted data to a DOCX file.

    Parameters:
        ocr_data (dict): Contains 'top_text' (str) and 'table' (list of lists or None).
        output_path (str): Path to save the DOCX file.
    """
    doc = Document()

    # === Set default font ===
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Add heading ===
    filename = os.path.basename(output_path)
    doc.add_heading(f"OCR Extract: {filename}", level=1)

    # === Add top paragraph text ===
    if ocr_data.get("top_text"):
        doc.add_paragraph(ocr_data["top_text"])

    # === Write table if present ===
    if ocr_data.get("table"):
        rows = ocr_data["table"]
        max_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cleaned_text = clean_cell_text(cell_text)

                if not cleaned_text and cell_text:
                    print(f"Filtered cell ({i},{j}): '{cell_text}'")

                cell = table.cell(i, j)
                cell.text = cleaned_text

                # === Right-align numeric-looking cells ===
                if cleaned_text and cleaned_text.replace(",", "").replace(".", "").isdigit():
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 2  # RIGHT

    # === Add timestamp footer (optional) ===
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Intense Quote')

    # === Ensure output folder exists and save ===
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    print("Saving to absolute path:", os.path.abspath(output_path))
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")


def batch_ocr_to_docx(input_folder, output_folder, ocr_function):
    os.makedirs(output_folder, exist_ok=True)
    for filename in os.listdir(input_folder):
        if filename.lower().endswith(".png"):
            input_path = os.path.join(input_folder, filename)
            image = Image.open(input_path).convert("RGB")
            ocr_data = ocr_function(image)
            output_filename = filename.rsplit(".", 1)[0] + ".docx"
            output_path = os.path.join(output_folder, output_filename)
            write_ocr_output_to_docx(ocr_data, output_path)

# ✅ This block must be at the bottom, outside of all functions
if __name__ == "__main__":
    from png_ocr import extract_structured_data

    input_image_path = "output_pages/page1.png"
    output_docx_path = "output_docx/page1.docx"
    debug_id = "page1"

    if not os.path.exists(input_image_path):
        print("❌ Image not found:", input_image_path)
        exit()

    print("🔍 Loading image...")
    image = Image.open(input_image_path).convert("RGB")

    print("🧠 Running OCR...")
    ocr_data = extract_structured_data(image, debug_id=debug_id)
    print("📝 OCR data received. Writing DOCX...")

    write_ocr_output_to_docx(ocr_data, output_docx_path)
